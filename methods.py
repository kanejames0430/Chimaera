#Imports:
import os
from nltk.stem import SnowballStemmer
import string
import re
import numpy as np
import numpy as np
import plotly.graph_objs as go
import fitz
from docx import Document
import win32com.client

def readData(dir):
    '''
    Given a directory of txt files containing the 1000 most common words of a language, read the contents of each
    text file into their respective lists and add it to a dictionary with the name of the language as the key, the
    list of words being the value. Text files should be named "language.txt" and their contents should only have 
    one word on a line with no other punctuation.
    '''
    languageDict = {}
    
    for file in (os.listdir(dir)):
        fileString = dir + '/' + file
        with open(fileString,"r",encoding='utf-8') as f:
            templist = f.readlines()
        languageList = []
        for x in templist:
            languageList.append(x.strip('\n'))
        languageDict[file.split("/")[-1].strip('.txt')] = languageList

    return languageDict

def stem(dictionary):
    '''
    tokenize takes in the dictionary of languages we create in readData(), and returns a dictionary of tokenized words.
    We use the snowball stemmer method because it covers a total of 17 languages, which we can implement later.
    '''
    dict = {}

    #for each language in the dictionary, 
    for j,k in dictionary.items():

        # initialize/reset the list of words
        tokList = []

        # Instantiate stemming class
        stemmer = SnowballStemmer(j)

        # For each word, we tokenize (find its stem) and append it to the list to retur
        for word in k:
            stem_word = stemmer.stem(word)  # stemming word
            tokList.append(stem_word)  # append to the list
        dict[j] = tokList
    return dict

def readRawText(path):
    '''
    Preprocessing: read the document and extract just the words into a list. Remove punctuation, digits, double spaces, new line/tab indicators.
    '''
    with open(path, 'r', encoding = 'utf-8') as f:
        content = (f.read())

        # string of items we dont want in our content
        removeItems = string.punctuation + string.digits + '\n' + '\t'

        # remove unwanted characters
        content = content.translate(str.maketrans('','',removeItems))

        # removes double spaces that are artifacts from removing numbers 
        # (e.g. ...over 200 million... -> ...over  million... which has a double space)
        content = re.sub(' +',' ', content)

        # change to all lowercase
        content = content.lower()

    return content

def freqsCalculator(content, tokenizedDictionary):
    '''
    from the tokenized dictionary: for every language, 
        and for every tokenized word,
            count the amount of times it appears in the document.
    content pertains to the processed text of the document. tokenizedDictionary is our processed dictionary
    '''
    content = content.split(' ')
    # reorganize tokDict into a dictionary of format {[language, tokenized word] : 0}
    freqs = {}

    # create a new dictionary with lists as entries in the words list instead of just the words
    # e.g. [word_1, word_2, ..., word_n] -> [(word_1, 0), (word_2,0), ..., (word_n, 0)]
    for x,y in tokenizedDictionary.items():
        for word in y:
            freqs[x,word] = 0

    for i,j in freqs.items():
        for word2 in content:
            if i[1] in word2:
                freqs[i] += 1        
    return freqs

def langScore(content, freqsDictionary):
    '''
    Takes in the frequency dictionary and sums over the scores for each language. 
    The content is the processed text from a document, which we will need the length of
    '''
    #initialize
    language_scores = {}

    #for each language, sum over the frequencies of each word respective to the language
    for (language, word), score in freqsDictionary.items():
        if language not in language_scores:
            language_scores[language] = score
        else:
            language_scores[language] += score
    # for each score, divide by the length of the content to normalize
    for lang, score2 in language_scores.items():
        finalScore = score2 / len(content)
        language_scores[lang] = finalScore
    return language_scores

def matrixBuilder(dict):
    '''
    Creates a matrix from the dictionary of language scores or list of vectors
    '''

    matrix = []
    labels = []

    for key, vec in dict.items():
        matrix.append(vec)
        labels.append(key)

    return np.array(matrix), labels


def project_to_2d(matrix):
    '''
    Project the matrix onto 2D space using Singular Value Decomposition (SVD).
    '''
    U, _, _ = np.linalg.svd(matrix)
    projected_matrix = U[:, :2]
    return projected_matrix

def project_to_3d(matrix):
    '''
    Project the matrix onto 3D space using Singular Value Decomposition (SVD).
    '''
    U, _, _ = np.linalg.svd(matrix)
    projected_matrix = U[:, :3]
    return projected_matrix

def calculate_distance_from_origin(points):
    '''
    For a list of points, calculate the norm, or the distance from the origin
    '''
    return np.linalg.norm(points, axis=1)

def normalize(array):
    '''
    Normalize the array to a range between 0 and 1.
    '''
    min_val = np.min(array)
    max_val = np.max(array)
    return (array - min_val) / (max_val - min_val)

def plot_2d_projection(matrix, labels):
    # Normalize each principal component
    norm_pc1 = normalize(matrix[:, 0])
    norm_pc2 = normalize(matrix[:, 1])

    # Combine the normalized principal components into RGB colors
    blue_channel = 128  # Set a constant blue channel value for visibility
    colors = np.stack([norm_pc1, norm_pc2, np.full_like(norm_pc1, blue_channel / 255)], axis=1)
    colors = (colors * 255).astype(int)  # Convert to RGB values in the range [0, 255]
    colors_hex = ['rgb({}, {}, {})'.format(r, g, b) for r, g, b in colors]

    trace = go.Scatter(
        x=matrix[:, 0],
        y=matrix[:, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=colors_hex,
            opacity=0.8,
        ),
        text=labels,
        hoverinfo='text'
    )

    layout = go.Layout(
        title='2D Projection',
        xaxis=dict(title='Principal Component 1'),
        yaxis=dict(title='Principal Component 2')
    )

    fig = go.Figure(data=[trace], layout=layout)
    fig.show()

def plot_3d_projection(matrix, labels):
    # Normalize each principal component
    norm_pc1 = normalize(matrix[:, 0])
    norm_pc2 = normalize(matrix[:, 1])
    norm_pc3 = normalize(matrix[:, 2])

    # Combine the normalized principal components into RGB colors
    colors = np.stack([norm_pc1, norm_pc2, norm_pc3], axis=1)
    colors = (colors * 255).astype(int)  # Convert to RGB values in the range [0, 255]
    colors_hex = ['rgb({}, {}, {})'.format(r, g, b) for r, g, b in colors]

    trace = go.Scatter3d(
        x=matrix[:, 0],
        y=matrix[:, 1],
        z=matrix[:, 2],
        mode='markers',
        marker=dict(
            size=10,
            color=colors_hex,
            opacity=0.8,
        ),
        text=labels,
        hoverinfo='text'
    )

    layout = go.Layout(
        title='3D Projection',
        scene=dict(
            xaxis=dict(title='Principal Component 1'),
            yaxis=dict(title='Principal Component 2'),
            zaxis=dict(title='Principal Component 3')
        )
    )

    fig = go.Figure(data=[trace], layout=layout)
    fig.show()

def extractText(path):
    '''
    From a given path, extract the contents of the file and return just the text.
    '''
    # gets the extention of the file
    extension = path.split('.')[-1]
    content = ''

    # Cases for different media files
    match extension:
        case 'pdf':
            # use fits to read raw text
            doc = fitz.open(path) 
            for page in doc: 
                content += page.get_text()

        case 'docx':
            # use pydocx to read raw text
            doc = Document(path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            content =  '\n'.join(full_text)

        case 'doc':
            # convert the .doc file and follow same procedure as .docx
            ''' 
            path = convert_doc_to_docx(path)
            doc = Document(path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            content =  '\n'.join(full_text)
            '''
            print("Feature not yet implemented!")
        case defualt:
            #the file is a txt file and no preprocessing needs to be handled
            with open(path,'r',encoding='utf-8') as f:
                content = f.read()
    return content,path,extension

def writeContent(content,path,extension,outDir):
    '''
    Takes in outputs from extract text and writes it to an appropriately names .txt file. The outDir is where we want
    to write the text to.
    '''
    # Remove the extension from the path string and add txt because we want to write the content to a text file
    fileName = path.split('/')[-1]
    outPath = outDir + '/' + fileName.replace(extension,'txt')
    with open(outPath,"w+",encoding='utf-8') as f:
        f.write(content)

def max_column_or_label(row):
    ''' 
    Checks the sum of the language score vector (in this case). If it is zero, then
    there may not be enough data collected  or the doc is in a language that is not recognized
    in our program.
    '''
    if row.sum() == 0:
        return "Not Enough Data"
    else:
        return row.idxmax()

def convert_doc_to_docx(file_path):
    ''' 
    Converts .doc files to .docx
    '''
    # Initialize the Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    # Open the .doc file
    doc = word.Documents.Open(file_path)

    # Define the new .docx file path
    new_file_path = os.path.splitext(file_path)[0] + ".docx"

    # Save the file as .docx
    doc.SaveAs(new_file_path, FileFormat=16)  # 16 represents the .docx format

    # Close the original document and quit Word
    doc.Close()
    word.Quit()

    return new_file_path